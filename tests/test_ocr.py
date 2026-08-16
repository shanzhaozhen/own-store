"""OCR 版面重建的测试。

刻意**不跑真实模型**：版面重建是纯几何逻辑，直接喂检测框就能测，
跑得快、不依赖模型文件、结果确定。真实引擎的联调用 needs_samples 标记的
那个用例（默认跳过）。
"""

from __future__ import annotations

import numpy as np
import pytest

from shop_print.core import ocr

PAGE_W = 1000
CHAR_H = 20.0


def line(
    text: str, x0: float, y0: float, *, char_w: float = 20.0, h: float = CHAR_H, score: float = 0.99
) -> ocr.OcrLine:
    return ocr.OcrLine(text=text, score=score, x0=x0, y0=y0, x1=x0 + char_w * len(text), y1=y0 + h)


def group(lines: list[ocr.OcrLine]) -> list[ocr.OcrParagraph]:
    rows = [ocr._merge_row(r) for r in ocr._cluster_rows(lines)]  # noqa: SLF001
    rows.sort(key=lambda r: r.y0)
    return ocr._group_paragraphs(rows, PAGE_W)  # noqa: SLF001


def test_同一行的多个框会被拼起来() -> None:
    lines = [line("乙方：李四", 400, 100), line("甲方：张三", 100, 102)]
    rows = ocr._cluster_rows(lines)  # noqa: SLF001
    assert len(rows) == 1
    merged = ocr._merge_row(rows[0])  # noqa: SLF001
    assert merged.text.startswith("甲方：张三")
    assert "乙方：李四" in merged.text


def test_行内间隙大会补空格() -> None:
    """不补空格的话"姓名张三"会连成一团，看不出是两栏。"""
    merged = ocr._merge_row([line("姓名", 100, 100), line("张三", 400, 100)])  # noqa: SLF001
    assert " " in merged.text


def test_不同行不会被聚到一起() -> None:
    lines = [line("第一行", 100, 100), line("第二行", 100, 140)]
    assert len(ocr._cluster_rows(lines)) == 2  # noqa: SLF001


def test_大标题不会把行阈值带偏() -> None:
    """阈值用中位字高而不是平均值 —— 一个大标题就能把平均值拉高。"""
    lines = [
        line("很大的标题", 300, 60, char_w=48, h=60),
        line("正文第一行", 100, 200),
        line("正文第二行", 100, 232),
    ]
    assert len(ocr._cluster_rows(lines)) == 3  # noqa: SLF001


def test_居中的大字判为标题() -> None:
    title_w = 6 * 40.0
    lines = [
        ocr.OcrLine(
            "房屋租赁合同", 0.99, (PAGE_W - title_w) / 2, 60, (PAGE_W + title_w) / 2, 60 + 34
        ),
        line("正文一行", 100, 200),
    ]
    paragraphs = group(lines)
    assert paragraphs[0].is_heading is True
    assert paragraphs[1].is_heading is False


def test_靠左的短行不会被误判为标题() -> None:
    paragraphs = group([line("甲方：张三", 100, 100), line("乙方：李四", 100, 132)])
    assert all(not p.is_heading for p in paragraphs)


def test_等距的独立短行要分成不同段() -> None:
    """证明、合同里常见"甲方 / 乙方"这种等距短行，光看行间距会被粘成一坨。"""
    lines = [
        line("甲方：张三", 100, 100),
        line("乙方：李四", 100, 132),
        line("丙方：王五", 100, 164),
    ]
    assert len(group(lines)) == 3


def test_满宽的连续行要合成一段() -> None:
    full = "满" * 40  # 写到右边界
    lines = [
        line(full, 100, 100),
        line(full, 100, 132),
        line("最后一行短的。", 100, 164),
    ]
    paragraphs = group(lines)
    assert len(paragraphs) == 1
    assert paragraphs[0].text == full + full + "最后一行短的。"


def test_首行缩进另起一段() -> None:
    full = "满" * 40
    lines = [
        line(full, 100, 100),
        line(full, 100, 132),
        line("　　新的一段从缩进开始" + "满" * 28, 140, 164),
    ]
    paragraphs = group(lines)
    assert len(paragraphs) == 2


def test_行间距明显变大也换段() -> None:
    full = "满" * 40
    lines = [line(full, 100, 100), line(full, 100, 132), line(full, 100, 260)]
    assert len(group(lines)) == 2


def test_窄文字块的每一行都独立成段() -> None:
    """整块文字宽度不到页宽 45% 时，没有任何一行能暴露真实版心，
    只能认为这些是独立短行（表单、落款），一行一段。"""
    lines = [line("姓名王五", 100, 100), line("性别男", 100, 132), line("年龄六十", 100, 164)]
    assert len(group(lines)) == 3


def test_正常宽度文字块按右边界判段末() -> None:
    """反过来：文字块占了页宽的大部分时，用实际最右那一行当右边界，
    没写满的行才算段末。这条和上一条配对，别只改一边。"""
    full = "满" * 35  # 700px 宽，占页宽 70%
    lines = [
        line(full, 100, 100),
        line(full, 100, 132),
        line("短短的收尾。", 100, 164),
        line(full, 100, 196),
    ]
    paragraphs = group(lines)
    assert len(paragraphs) == 2
    assert paragraphs[0].lines[-1].text == "短短的收尾。"


def test_低置信度的行会被挑出来提示核对() -> None:
    result = ocr.OcrResult(
        lines=[line("清楚的字", 100, 100), line("模糊的字", 100, 140, score=0.42)],
    )
    assert [line_.text for line_ in result.low_confidence_lines] == ["模糊的字"]


def test_没识别出东西时是空结果() -> None:
    assert ocr.OcrResult().is_empty is True
    assert ocr.OcrResult().text == ""


def test_导出docx和txt(tmp_path) -> None:
    result = ocr.OcrResult(
        paragraphs=[
            ocr.OcrParagraph(lines=[line("房屋租赁合同", 300, 60)], is_heading=True, centered=True),
            ocr.OcrParagraph(lines=[line("正文内容一段。", 100, 200)]),
        ],
        lines=[],
        page_width=PAGE_W,
        page_height=1400,
    )
    docx_path = ocr.to_docx(result, tmp_path / "out.docx")
    txt_path = ocr.to_txt(result, tmp_path / "out.txt")
    assert docx_path.exists() and docx_path.stat().st_size > 0
    assert txt_path.read_text(encoding="utf-8") == "房屋租赁合同\n正文内容一段。"

    from docx import Document

    document = Document(docx_path)
    assert [p.text for p in document.paragraphs] == ["房屋租赁合同", "正文内容一段。"]
    assert document.paragraphs[0].runs[0].bold is True


@pytest.mark.needs_samples
def test_真实引擎联调(tmp_path) -> None:
    """跑真模型。需要 assets/models 里有模型文件，默认不在 CI 里跑。

    pytest -m needs_samples
    """
    import pymupdf

    document = pymupdf.open()
    page = document.new_page(width=595, height=842)
    font = pymupdf.Font("china-s")
    page.insert_font(fontname="cjk", fontbuffer=font.buffer)
    page.insert_text((70, 120), "房屋租赁合同", fontname="cjk", fontsize=18)
    image_path = tmp_path / "sample.png"
    page.get_pixmap(dpi=200).save(image_path)

    result = ocr.recognize(np.asarray(ocr.enhance_mod.load_image(image_path)))
    assert "租赁" in result.text
