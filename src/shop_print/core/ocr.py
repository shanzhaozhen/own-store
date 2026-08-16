"""照片转文字文档：OCR + 版面重建 → 可编辑的 Word。

需求原话是「拍照的图片需要转成文字文档并且排版」—— 重点在**并且排版**。
交付的不是一堆文字，是一份能直接改的 Word。

引擎用本地 RapidOCR（onnxruntime CPU），模型随包分发、运行时不联网。
为什么不是 PaddleOCR：见 docs/decisions/ADR-002-OCR引擎选择.md。
版面重建的规则见 docs/05-OCR与版面重建.md。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Any

import cv2
import numpy as np

from .. import paths
from ..texts import ErrorKind
from . import enhance as enhance_mod
from .errors import ShopPrintError

logger = logging.getLogger(__name__)

# 低于这个置信度就在界面上标出来，提示长辈核对
LOW_CONFIDENCE = 0.70

# 输出的 Word 用 A4：店里只有 A4/A3，和 convert / cards 出的 PDF 保持同一个物理尺寸
A4_WIDTH_MM = 210.0
A4_HEIGHT_MM = 297.0

# 行聚类：两个框的 y 中心差小于「中位字高 × 这个系数」就算同一行
_ROW_TOLERANCE = 0.6
# 段落切分：行间距大于「中位行间距 × 这个系数」就算换段
_PARAGRAPH_GAP = 1.8
# 上一行右端比右边界短这么多个字宽，就认为它是段末
_SHORT_LINE_SLACK_CHARS = 2.0
# 这一行左端比左边界缩进这么多个字宽，就认为是首行缩进、另起一段
_INDENT_SLACK_CHARS = 1.5
# 相邻两行字高比例超过这个值，就认为不是同一个文本块（标题 vs 正文）
_FONT_SIZE_JUMP = 1.4
# 整页文字块宽度不到页宽的这个比例，就认为这些是独立短行而不是折行的段落
_MIN_BLOCK_WIDTH_RATIO = 0.45
# 居中判定
_CENTERED_MAX_WIDTH_RATIO = 0.60
_CENTERED_MAX_OFFSET_RATIO = 0.08
# 标题判定：字高明显大于中位字高
_HEADING_HEIGHT_RATIO = 1.3


@dataclass
class OcrLine:
    """识别出的一行（可能由多个检测框拼成）。"""

    text: str
    score: float
    x0: float
    y0: float
    x1: float
    y1: float

    @property
    def width(self) -> float:
        return self.x1 - self.x0

    @property
    def height(self) -> float:
        return self.y1 - self.y0

    @property
    def y_center(self) -> float:
        return (self.y0 + self.y1) / 2.0

    @property
    def x_center(self) -> float:
        return (self.x0 + self.x1) / 2.0


@dataclass
class OcrParagraph:
    lines: list[OcrLine] = field(default_factory=list)
    is_heading: bool = False
    centered: bool = False

    @property
    def text(self) -> str:
        return "".join(line.text for line in self.lines)

    @property
    def min_score(self) -> float:
        return min((line.score for line in self.lines), default=0.0)


@dataclass
class OcrResult:
    paragraphs: list[OcrParagraph] = field(default_factory=list)
    lines: list[OcrLine] = field(default_factory=list)
    page_width: int = 0
    page_height: int = 0

    @property
    def text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs)

    @property
    def is_empty(self) -> bool:
        return not self.lines

    @property
    def low_confidence_lines(self) -> list[OcrLine]:
        """置信度低的行。界面上标出来，提示长辈核对 —— 识别不可能 100% 准，
        假装完美只会让顾客拿到错的文档。"""
        return [line for line in self.lines if line.score < LOW_CONFIDENCE]


@lru_cache(maxsize=1)
def get_engine() -> Any:
    """RapidOCR 引擎（进程内单例）。初始化要几百毫秒，不能每次识别都建。"""
    try:
        from rapidocr import RapidOCR
    except ImportError as exc:  # pragma: no cover
        raise ShopPrintError(ErrorKind.UNKNOWN, f"OCR 引擎装不上：{exc}") from exc

    model_dir = paths.models_dir()
    model_dir.mkdir(parents=True, exist_ok=True)
    # model_root_dir 指到随包的模型目录：店铺网络不一定稳，
    # 首次使用不能卡在"正在下载模型"上。
    return RapidOCR(
        params={
            "Global.model_root_dir": str(model_dir),
            "Global.log_level": "warning",
        }
    )


def _to_bgr(image: np.ndarray) -> np.ndarray:
    if image.ndim == 2:
        return cv2.cvtColor(image, cv2.COLOR_GRAY2BGR)
    return image


def _boxes_to_lines(boxes: Any, txts: Any, scores: Any) -> list[OcrLine]:
    lines: list[OcrLine] = []
    for box, text, score in zip(boxes, txts, scores, strict=False):
        if not text:
            continue
        points = np.asarray(box, dtype=np.float32).reshape(-1, 2)
        lines.append(
            OcrLine(
                text=str(text),
                score=float(score),
                x0=float(points[:, 0].min()),
                y0=float(points[:, 1].min()),
                x1=float(points[:, 0].max()),
                y1=float(points[:, 1].max()),
            )
        )
    return lines


def _median(values: list[float], fallback: float) -> float:
    return float(np.median(values)) if values else fallback


def _cluster_rows(boxes: list[OcrLine]) -> list[list[OcrLine]]:
    """按 y 中心把检测框聚成行。

    阈值用**中位**字高而不是平均值：一个大标题就能把平均值带偏。
    """
    if not boxes:
        return []
    median_h = _median([b.height for b in boxes], 1.0)
    tolerance = max(1.0, median_h * _ROW_TOLERANCE)

    rows: list[list[OcrLine]] = []
    current: list[OcrLine] = []
    current_center = 0.0
    for box in sorted(boxes, key=lambda b: b.y_center):
        if current and abs(box.y_center - current_center) > tolerance:
            rows.append(current)
            current = []
        current.append(box)
        current_center = float(np.mean([b.y_center for b in current]))
    if current:
        rows.append(current)
    return rows


def _merge_row(row: list[OcrLine]) -> OcrLine:
    """把一行里的多个框按 x 排序拼成一行文字。

    中文之间不加空格；水平间隙超过一个字宽时补一个空格 —— 那通常是
    分栏或制表位，不补的话"姓名张三"会连成一团。
    """
    row = sorted(row, key=lambda b: b.x0)
    char_widths = [b.width / max(1, len(b.text)) for b in row]
    char_width = _median(char_widths, 10.0)

    parts: list[str] = []
    previous: OcrLine | None = None
    for box in row:
        if previous is not None and (box.x0 - previous.x1) > char_width:
            parts.append(" ")
        parts.append(box.text)
        previous = box
    return OcrLine(
        text="".join(parts),
        score=min(b.score for b in row),
        x0=min(b.x0 for b in row),
        y0=min(b.y0 for b in row),
        x1=max(b.x1 for b in row),
        y1=max(b.y1 for b in row),
    )


def _is_centered(line: OcrLine, page_width: int) -> bool:
    if page_width <= 0:
        return False
    narrow = line.width < page_width * _CENTERED_MAX_WIDTH_RATIO
    near_center = abs(line.x_center - page_width / 2.0) < page_width * _CENTERED_MAX_OFFSET_RATIO
    return narrow and near_center


def _median_height_excluding(heights: list[float], index: int) -> float:
    """除掉第 index 行之后的中位字高。

    判断"这行是不是标题"要和**别的行**比。整页只有两三行时，把自己也算进
    中位数会把基准抬高，大标题反而判不出来。
    """
    others = heights[:index] + heights[index + 1 :]
    return _median(others, heights[index] if heights else 1.0)


def _group_paragraphs(rows: list[OcrLine], page_width: int) -> list[OcrParagraph]:
    """行 → 段落。

    只看行间距不够用，再加三条排版常识：

    - **上一行明显没写到右边界** → 它是段末，下一行另起一段。
      （证明、合同里常见"甲方：张三 / 乙方：李四"这种等距独立短行，
      光看间距会被粘成一坨。）
    - **这一行明显往右缩进** → 首行缩进，另起一段
    - **两行字号差得多** → 不是同一个文本块（标题 vs 正文）

    右边界怎么估：正常情况用**实际最右**的那一行。但整页文字块本身就很窄
    （宽度不到页宽的 45%）时，没有任何一行能暴露真实版心 —— 这种情况下
    这些行本来就是独立短行而不是折行的段落，改用"按左边距镜像"的右边界，
    让每一行都判成段末。宁可切多也不要粘连：一行一段视觉上仍接近原件，
    粘成一坨就明显是错的。
    """
    if not rows:
        return []
    heights = [r.height for r in rows]
    median_h = _median(heights, 1.0)
    char_width = _median([r.width / max(1, len(r.text)) for r in rows], median_h)

    gaps = [rows[i].y0 - rows[i - 1].y1 for i in range(1, len(rows))]
    positive = [g for g in gaps if g > 0]
    # 用 25 分位而不是中位数估"正常行距"：段内行距占多数且偏小，
    # 中位数会被少数换段的大间距抬起来，导致该切的地方切不开。
    leading = float(np.percentile(positive, 25)) if positive else median_h * 0.4
    gap_threshold = max(median_h * 0.8, leading * _PARAGRAPH_GAP)

    text_left = float(np.percentile([r.x0 for r in rows], 10))
    observed_right = max(r.x1 for r in rows)
    narrow_block = (observed_right - text_left) < page_width * _MIN_BLOCK_WIDTH_RATIO
    text_right = (page_width - text_left) if narrow_block else observed_right
    short_line_slack = char_width * _SHORT_LINE_SLACK_CHARS
    indent_slack = char_width * _INDENT_SLACK_CHARS

    paragraphs: list[OcrParagraph] = []
    current = OcrParagraph(lines=[rows[0]])
    for index in range(1, len(rows)):
        previous, line = rows[index - 1], rows[index]
        taller, shorter = max(previous.height, line.height), min(previous.height, line.height)
        if (
            (line.y0 - previous.y1) > gap_threshold
            or previous.x1 < text_right - short_line_slack
            or line.x0 > text_left + indent_slack
            or (shorter > 0 and taller / shorter > _FONT_SIZE_JUMP)
        ):
            paragraphs.append(current)
            current = OcrParagraph(lines=[line])
        else:
            current.lines.append(line)
    paragraphs.append(current)

    index_of = {id(row): i for i, row in enumerate(rows)}
    for paragraph in paragraphs:
        first = paragraph.lines[0]
        paragraph.centered = len(paragraph.lines) == 1 and _is_centered(first, page_width)
        baseline = _median_height_excluding(heights, index_of[id(first)])
        paragraph.is_heading = (
            paragraph.centered and first.height > baseline * _HEADING_HEIGHT_RATIO
        )
    return paragraphs


def recognize(image: np.ndarray, preprocess: bool = True) -> OcrResult:
    """识别一张图。

    默认先过 enhance.prepare_for_ocr（裁正 + 去阴影 + 灰度）：文字行变水平让
    检测框更准，暗部文字不被吞掉。**不做二值化** —— 过度二值化吃掉笔画细节，
    反而降低识别率。
    """
    prepared = enhance_mod.prepare_for_ocr(image) if preprocess else image
    height, width = prepared.shape[:2]

    engine = get_engine()
    try:
        raw = engine(_to_bgr(prepared))
    except Exception as exc:
        logger.exception("OCR 失败")
        raise ShopPrintError(ErrorKind.OCR_EMPTY, f"识别出错：{exc}") from exc

    boxes = getattr(raw, "boxes", None)
    txts = getattr(raw, "txts", None)
    scores = getattr(raw, "scores", None)
    if boxes is None or txts is None or not len(txts):
        return OcrResult(page_width=width, page_height=height)

    detections = _boxes_to_lines(boxes, txts, scores if scores is not None else [1.0] * len(txts))
    rows = [_merge_row(row) for row in _cluster_rows(detections)]
    rows.sort(key=lambda r: r.y0)
    return OcrResult(
        paragraphs=_group_paragraphs(rows, width),
        lines=rows,
        page_width=width,
        page_height=height,
    )


def recognize_file(path: str | Path, preprocess: bool = True) -> OcrResult:
    return recognize(enhance_mod.load_image(path), preprocess=preprocess)


def _set_cjk_font(run: Any, name: str = "宋体") -> None:
    """python-docx 设中文字体必须同时写 w:eastAsia，只设 name 中文会退回默认字体。"""
    from docx.oxml.ns import qn

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)  # noqa: SLF001 —— python-docx 官方写法


def to_docx(result: OcrResult, out_path: str | Path) -> Path:
    """生成可编辑的 Word。**A4 纸**、正文宋体小四、行距 1.5、首行缩进 2 字符。

    页面尺寸必须显式设成 A4：python-docx 自带的模板是美国 Letter
    （215.9×279.4mm），店里只有 A4/A3 纸，拿 Letter 的文档去打，Word 会
    自己缩放或者把版面挪位 —— 和 PDF 那条路出来的尺寸也就不一致了。
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    for section in document.sections:
        section.page_width = Mm(A4_WIDTH_MM)
        section.page_height = Mm(A4_HEIGHT_MM)
        section.top_margin = section.bottom_margin = Pt(72)
        section.left_margin = section.right_margin = Pt(80)

    for paragraph in result.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        block = document.add_paragraph()
        fmt = block.paragraph_format
        fmt.line_spacing = 1.5
        fmt.space_after = Pt(0)

        run = block.add_run(text)
        if paragraph.is_heading:
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(12)
            run.bold = True
            run.font.size = Pt(16)
        elif paragraph.centered:
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(12)
        else:
            fmt.first_line_indent = Pt(24)  # 12pt 的两个字
            run.font.size = Pt(12)
        _set_cjk_font(run)

    document.save(out_path)
    return out_path


def to_txt(result: OcrResult, out_path: str | Path) -> Path:
    """另存一份纯文本，方便长辈直接复制粘贴。"""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(result.text, encoding="utf-8")
    return out_path


# ── 保留原有排版的 Word ─────────────────────────────────────────
_W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
_VML_NS = 'xmlns:v="urn:schemas-microsoft-com:vml"'
# 文本框比文字本身放大一点，免得 Word 按自己的字宽算完就折行了
_BOX_SLACK = 1.18
# 检测框的高度包含上下一点余量，字号按这个系数折算
_FONT_FROM_BOX = 0.78
_MIN_FONT_PT = 6.0
_PT_PER_MM = 72.0 / 25.4


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    )


def _textbox_xml(
    lines: list[str], x_pt: float, y_pt: float, w_pt: float, h_pt: float, font_pt: float, index: int
) -> str:
    """一个绝对定位的文本框（VML）。**框里的字仍然能选中、复制、修改。**

    用 VML（`v:rect` + `v:textbox`）而不是新式 DrawingML：写法短得多，
    Word 2007 以来一直支持。位置相对**页面**，这样才能和原图坐标一一对应。
    """
    half = max(2, round(font_pt * 2))  # w:sz 的单位是半磅
    body = "".join(
        '<w:p><w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f'<w:r><w:rPr><w:rFonts w:ascii="宋体" w:eastAsia="宋体" w:hAnsi="宋体"/>'
        f'<w:sz w:val="{half}"/><w:szCs w:val="{half}"/></w:rPr>'
        f'<w:t xml:space="preserve">{_escape(line)}</w:t></w:r></w:p>'
        for line in lines
    )
    style = (
        f"position:absolute;margin-left:{x_pt:.2f}pt;margin-top:{y_pt:.2f}pt;"
        f"width:{w_pt:.2f}pt;height:{h_pt:.2f}pt;z-index:{index};"
        "mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
    )
    return (
        f"<w:r><w:pict>"
        f'<v:rect style="{style}" filled="f" stroked="f">'
        f'<v:textbox inset="0,0,0,0"><w:txbxContent>{body}</w:txbxContent></v:textbox>'
        "</v:rect></w:pict></w:r>"
    )


def to_docx_layout(
    result: OcrResult,
    out_path: str | Path,
    page_width_mm: float = A4_WIDTH_MM,
    page_height_mm: float = A4_HEIGHT_MM,
) -> Path:
    """**保留原有排版**的 Word：每段文字放进一个按原位置摆好的文本框。

    和 `to_docx()`（顺排成普通段落）的区别：这里照着照片上的坐标摆，表格、
    多栏、签名的位置都还在原处；文字仍在文本框里，能选中、复制、改错字。
    代价是它更像"照着原件重排的版"，不方便整篇重新排版 —— 那种需求用 `to_docx()`。

    坐标怎么换：识别是在整页图上做的，所以 `1 像素 = 页宽 / 图宽`，再换成磅。
    """
    from docx import Document
    from docx.oxml import parse_xml
    from docx.shared import Mm, Pt

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if result.page_width <= 0:
        raise ShopPrintError(ErrorKind.OCR_EMPTY, "没有页面尺寸，排不出版")

    mm_per_px = page_width_mm / float(result.page_width)

    document = Document()
    for section in document.sections:
        section.page_width = Mm(page_width_mm)
        section.page_height = Mm(page_height_mm)
        # 文本框是相对页面定位的，页边距只影响那个空的锚段落，设小一点省地方
        section.top_margin = section.bottom_margin = Mm(10)
        section.left_margin = section.right_margin = Mm(10)

    # 文本框必须挂在某个段落上（Word 的要求），所以留一个空段落当锚点
    anchor = document.add_paragraph()
    anchor.paragraph_format.space_after = Pt(0)
    cursor = anchor._p  # noqa: SLF001 —— 往后依次插入，保持阅读顺序

    for index, paragraph in enumerate(result.paragraphs, start=1):
        lines = [line.text for line in paragraph.lines if line.text.strip()]
        if not lines:
            continue
        x0 = min(line.x0 for line in paragraph.lines)
        y0 = min(line.y0 for line in paragraph.lines)
        x1 = max(line.x1 for line in paragraph.lines)
        y1 = max(line.y1 for line in paragraph.lines)
        font_pt = max(
            _MIN_FONT_PT,
            _median([line.height for line in paragraph.lines], 1.0)
            * mm_per_px
            * _PT_PER_MM
            * _FONT_FROM_BOX,
        )
        width_pt = max(font_pt, (x1 - x0) * mm_per_px * _PT_PER_MM * _BOX_SLACK)
        height_pt = max(font_pt * 1.4, (y1 - y0) * mm_per_px * _PT_PER_MM * 1.4)
        xml = (
            f"<w:p {_W_NS} {_VML_NS}>"
            + _textbox_xml(
                lines,
                x0 * mm_per_px * _PT_PER_MM,
                y0 * mm_per_px * _PT_PER_MM,
                width_pt,
                height_pt,
                font_pt,
                index,
            )
            + "</w:p>"
        )
        # 按阅读顺序往后插：都插在锚点后面的话顺序会整个反过来，
        # 虽然位置一样，但在 Word 里按 Tab 或者用读屏软件时顺序是乱的
        element = parse_xml(xml)
        cursor.addnext(element)
        cursor = element

    document.save(out_path)
    return out_path


def image_to_document(src: str | Path, out_dir: Path | None = None) -> tuple[Path, Path, OcrResult]:
    """一步到底：图片 → (docx, txt, 识别结果)。"""
    src = Path(src)
    result = recognize_file(src)
    if result.is_empty:
        raise ShopPrintError(ErrorKind.OCR_EMPTY, f"{src.name} 没识别出文字")
    target = out_dir or paths.output_dir()
    target.mkdir(parents=True, exist_ok=True)
    docx_path = to_docx(result, target / f"{src.stem}.docx")
    txt_path = to_txt(result, target / f"{src.stem}.txt")
    return docx_path, txt_path, result
