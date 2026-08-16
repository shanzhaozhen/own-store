"""云端高精度识别的接口。**v1 只定义协议，不实现任何 provider。**

为什么留这个口子：本地 RapidOCR 对付常规文档够用，但**表格、多栏排版、
复杂公文吃力**。云端视觉大模型或带版面分析的云 OCR 能直接输出结构化文本，
这类场景质量高一个档次。

为什么统一返回 Markdown 而不是坐标框：换 provider（百度 / 腾讯 / 视觉大模型）
不用改下游，下游只需要一个 Markdown → docx 转换器。
见 docs/decisions/ADR-002-OCR引擎选择.md。

配置在 config.json 的 ocr 段；没填 provider 时界面上「高精度识别」按钮置灰。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Protocol, runtime_checkable

import numpy as np

from ..config import OcrConfig
from ..texts import ErrorKind
from .errors import ShopPrintError


@runtime_checkable
class CloudOcrProvider(Protocol):
    """云端识别 provider。实现方只要满足这个协议就能插进来。"""

    name: str

    def recognize(self, image: np.ndarray) -> str:
        """返回 Markdown 文本（标题、列表、表格）。"""
        ...


_REGISTRY: dict[str, CloudOcrProvider] = {}


def register(provider: CloudOcrProvider) -> None:
    _REGISTRY[provider.name] = provider


def available() -> list[str]:
    return sorted(_REGISTRY)


def is_configured(config: OcrConfig) -> bool:
    """界面用这个判断「高精度识别」按钮该不该置灰。"""
    return bool(config.cloud_provider) and config.cloud_provider in _REGISTRY


def recognize(image: np.ndarray, config: OcrConfig) -> str:
    if not config.cloud_provider:
        raise ShopPrintError(ErrorKind.UNKNOWN, "还没有配置云端识别")
    provider = _REGISTRY.get(config.cloud_provider)
    if provider is None:
        raise ShopPrintError(
            ErrorKind.UNKNOWN,
            f"没有注册名为「{config.cloud_provider}」的云端识别，已注册：{available()}",
        )
    return provider.recognize(image)


_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_LIST_ITEM = re.compile(r"^\s*([-*+]|\d+[.)])\s+(.*)$")
_TABLE_SEPARATOR = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def markdown_to_docx(markdown: str, out_path: str | Path) -> Path:
    """Markdown → Word。只处理标题、列表、表格、段落这几样，够用就行。"""
    from docx import Document
    from docx.shared import Pt

    from .ocr import _set_cjk_font

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        if not table_buffer:
            return
        columns = max(len(row) for row in table_buffer)
        table = document.add_table(rows=0, cols=columns)
        table.style = "Table Grid"
        for row in table_buffer:
            cells = table.add_row().cells
            for index in range(columns):
                cells[index].text = row[index] if index < len(row) else ""
        table_buffer.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("|") and not _TABLE_SEPARATOR.match(line):
            table_buffer.append([cell.strip() for cell in line.strip("|").split("|")])
            continue
        if _TABLE_SEPARATOR.match(line) and table_buffer:
            continue  # 表头分隔行，跳过
        flush_table()

        if not line.strip():
            continue
        heading = _HEADING.match(line)
        if heading:
            document.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 4))
            continue
        item = _LIST_ITEM.match(line)
        if item:
            style = "List Number" if item.group(1)[0].isdigit() else "List Bullet"
            block = document.add_paragraph(item.group(2).strip(), style=style)
            if block.runs:
                _set_cjk_font(block.runs[0])
            continue

        block = document.add_paragraph()
        block.paragraph_format.line_spacing = 1.5
        block.paragraph_format.first_line_indent = Pt(24)
        run = block.add_run(line.strip())
        run.font.size = Pt(12)
        _set_cjk_font(run)

    flush_table()
    document.save(out_path)
    return out_path
